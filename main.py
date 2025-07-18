#!/usr/bin/env python3
"""
FastAPI Zoom Meeting Transcript Generator
Upload audio files via POST request to get transcripts
"""

import os
import re
import tempfile
import shutil
import uuid
import asyncio
import time
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any
from contextlib import asynccontextmanager

import whisper
import numpy as np
from pydub import AudioSegment
from docx import Document

from fastapi import FastAPI, File, UploadFile, HTTPException, Request, Depends, Query
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.trustedhost import TrustedHostMiddleware
from starlette.background import BackgroundTask
from starlette.middleware.base import BaseHTTPMiddleware

# Local imports
from config import get_settings, Settings, is_allowed_file_extension, format_file_size
from logging_config import setup_logging, get_logger, log_request_start, log_request_end, log_error
from exceptions import (
    FileValidationError, AudioProcessingError, TranscriptionError, ModelLoadError,
    BadRequestException, InternalServerErrorException, ServiceUnavailableException,
    ErrorCodes, create_error_response
)
from models import (
    TranscriptionRequest, HealthCheckResponse, HealthStatus,
    TranscriptionResponse, ErrorResponse, FileUploadInfo
)
from middleware import (
    RateLimitMiddleware, SecurityHeadersMiddleware,
    RequestSizeLimitMiddleware, ErrorHandlingMiddleware
)

# Initialize logging
setup_logging()
logger = get_logger(__name__)

# Global variables
whisper_model = None
app_start_time = None


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Lifespan event handler for startup and shutdown"""
    global whisper_model, app_start_time

    # Startup
    app_start_time = time.time()
    settings = get_settings()

    logger.info("Starting Zoom Meeting Transcript API", extra={
        "version": settings.app_version,
        "whisper_model": settings.whisper_model_name,
        "device": settings.whisper_device
    })

    try:
        logger.info(f"Loading Whisper model: {settings.whisper_model_name}")
        whisper_model = whisper.load_model(
            settings.whisper_model_name,
            device=settings.whisper_device
        )
        logger.info("Whisper model loaded successfully")
    except Exception as e:
        logger.error(f"Failed to load Whisper model: {str(e)}", exc_info=True)
        raise ModelLoadError(f"Failed to load Whisper model: {str(e)}")

    # Create required directories
    os.makedirs(settings.upload_dir, exist_ok=True)
    if settings.temp_dir:
        os.makedirs(settings.temp_dir, exist_ok=True)

    logger.info("Application startup completed")

    yield

    # Shutdown
    logger.info("Shutting down application")
    # Cleanup resources if needed


def get_app_settings() -> Settings:
    """Dependency to get application settings"""
    return get_settings()


# Create FastAPI app
settings = get_settings()
app = FastAPI(
    title=settings.app_name,
    description=settings.app_description,
    version=settings.app_version,
    lifespan=lifespan,
    docs_url="/docs" if settings.debug else None,
    redoc_url="/redoc" if settings.debug else None,
    openapi_tags=[
        {
            "name": "health",
            "description": "Health check and system status endpoints"
        },
        {
            "name": "transcription",
            "description": "Audio transcription endpoints"
        }
    ],
    contact={
        "name": "API Support",
        "email": "support@example.com",
    },
    license_info={
        "name": "MIT",
        "url": "https://opensource.org/licenses/MIT",
    },
)

# Add middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=settings.cors_origins,
    allow_credentials=True,
    allow_methods=["GET", "POST"],
    allow_headers=["*"],
)

app.add_middleware(
    TrustedHostMiddleware,
    allowed_hosts=settings.allowed_hosts
)

# Add custom middleware (order matters - last added is executed first)
app.add_middleware(ErrorHandlingMiddleware)
app.add_middleware(SecurityHeadersMiddleware)
app.add_middleware(RequestSizeLimitMiddleware)
# Note: Rate limiting is handled by Vercel in production
if settings.debug:
    app.add_middleware(RateLimitMiddleware)


class RequestTrackingMiddleware(BaseHTTPMiddleware):
    """Middleware to track requests and add request IDs"""

    async def dispatch(self, request: Request, call_next):
        # Generate request ID
        request_id = str(uuid.uuid4())
        request.state.request_id = request_id

        # Log request start
        start_time = time.time()
        log_request_start(
            logger,
            request_id,
            str(request.url.path),
            request.method
        )

        # Process request
        try:
            response = await call_next(request)

            # Log successful completion
            duration = time.time() - start_time
            log_request_end(
                logger,
                request_id,
                str(request.url.path),
                request.method,
                response.status_code,
                duration
            )

            # Add request ID to response headers
            response.headers["X-Request-ID"] = request_id
            return response

        except Exception as e:
            # Log error
            duration = time.time() - start_time
            log_error(logger, request_id, e, "request_processing")

            # Re-raise the exception
            raise


# Add request tracking middleware
app.add_middleware(RequestTrackingMiddleware)


# Exception handlers
@app.exception_handler(FileValidationError)
async def file_validation_exception_handler(request: Request, exc: FileValidationError):
    """Handle file validation errors"""
    request_id = getattr(request.state, 'request_id', str(uuid.uuid4()))

    error_response = create_error_response(
        ErrorCodes.VALIDATION_ERROR,
        str(exc),
        details=exc.details,
        request_id=request_id
    )
    error_response["error"]["timestamp"] = datetime.now(datetime.timezone.utc).isoformat()

    return JSONResponse(
        status_code=400,
        content=error_response
    )


@app.exception_handler(AudioProcessingError)
async def audio_processing_exception_handler(request: Request, exc: AudioProcessingError):
    """Handle audio processing errors"""
    request_id = getattr(request.state, 'request_id', str(uuid.uuid4()))

    error_response = create_error_response(
        ErrorCodes.AUDIO_PROCESSING_FAILED,
        str(exc),
        details=exc.details,
        request_id=request_id
    )
    error_response["error"]["timestamp"] = datetime.now(datetime.timezone.utc).isoformat()

    return JSONResponse(
        status_code=422,
        content=error_response
    )


@app.exception_handler(TranscriptionError)
async def transcription_exception_handler(request: Request, exc: TranscriptionError):
    """Handle transcription errors"""
    request_id = getattr(request.state, 'request_id', str(uuid.uuid4()))

    error_response = create_error_response(
        ErrorCodes.TRANSCRIPTION_FAILED,
        str(exc),
        details=exc.details,
        request_id=request_id
    )
    error_response["error"]["timestamp"] = datetime.now(datetime.timezone.utc).isoformat()

    return JSONResponse(
        status_code=500,
        content=error_response
    )


@app.exception_handler(ModelLoadError)
async def model_load_exception_handler(request: Request, exc: ModelLoadError):
    """Handle model loading errors"""
    request_id = getattr(request.state, 'request_id', str(uuid.uuid4()))

    error_response = create_error_response(
        ErrorCodes.MODEL_NOT_LOADED,
        str(exc),
        details=exc.details,
        request_id=request_id
    )
    error_response["error"]["timestamp"] = datetime.now(datetime.timezone.utc).isoformat()

    return JSONResponse(
        status_code=503,
        content=error_response
    )


@app.get(
    "/",
    response_model=HealthCheckResponse,
    tags=["health"],
    summary="Health Check",
    description="Check the health and status of the transcription service",
    responses={
        200: {
            "description": "Service is healthy",
            "model": HealthCheckResponse
        },
        503: {
            "description": "Service is unhealthy",
            "model": ErrorResponse
        }
    }
)
async def health_check(settings: Settings = Depends(get_app_settings)):
    """
    Enhanced health check endpoint that provides detailed system status information.

    Returns information about:
    - Service status (healthy/unhealthy/degraded)
    - Whisper model loading status
    - System uptime
    - Disk space availability
    - Upload directory accessibility
    """
    global app_start_time

    # Calculate uptime
    uptime = time.time() - app_start_time if app_start_time else 0

    # Perform health checks
    checks = {}
    status = HealthStatus.HEALTHY

    # Check Whisper model
    if whisper_model is not None:
        checks["whisper_model"] = "loaded"
    else:
        checks["whisper_model"] = "not_loaded"
        status = HealthStatus.UNHEALTHY

    # Check disk space
    try:
        disk_usage = shutil.disk_usage(settings.upload_dir)
        free_gb = disk_usage.free / (1024**3)
        if free_gb < 1.0:  # Less than 1GB free
            checks["disk_space"] = f"low ({free_gb:.1f}GB free)"
            status = HealthStatus.DEGRADED
        else:
            checks["disk_space"] = f"sufficient ({free_gb:.1f}GB free)"
    except Exception:
        checks["disk_space"] = "unknown"
        status = HealthStatus.DEGRADED

    # Check upload directory
    if os.path.exists(settings.upload_dir) and os.access(settings.upload_dir, os.W_OK):
        checks["upload_directory"] = "accessible"
    else:
        checks["upload_directory"] = "inaccessible"
        status = HealthStatus.UNHEALTHY

    message = "Zoom Meeting Transcript API is running"
    if status == HealthStatus.UNHEALTHY:
        message = "Service is unhealthy"
    elif status == HealthStatus.DEGRADED:
        message = "Service is running with degraded performance"

    return HealthCheckResponse(
        status=status,
        message=message,
        timestamp=datetime.now(datetime.timezone.utc),
        version=settings.app_version,
        model_loaded=whisper_model is not None,
        uptime_seconds=uptime,
        checks=checks
    )


async def validate_uploaded_file(file: UploadFile, settings: Settings) -> FileUploadInfo:
    """Validate uploaded file"""
    if not file.filename:
        raise FileValidationError("No filename provided")

    # Check file extension
    file_extension = Path(file.filename).suffix.lower()
    if not is_allowed_file_extension(file.filename):
        raise FileValidationError(
            f"Unsupported file type '{file_extension}'. "
            f"Allowed: {', '.join(settings.allowed_extensions)}"
        )

    # Check file size
    file.file.seek(0, 2)  # Seek to end
    file_size = file.file.tell()
    file.file.seek(0)  # Reset to beginning

    if file_size == 0:
        raise FileValidationError("File is empty")

    if file_size > settings.max_file_size:
        raise FileValidationError(
            f"File too large ({format_file_size(file_size)}). "
            f"Maximum allowed: {format_file_size(settings.max_file_size)}"
        )

    return FileUploadInfo(
        filename=file.filename,
        size=file_size,
        content_type=file.content_type or "application/octet-stream",
        extension=file_extension
    )


@app.post(
    "/transcribe",
    response_class=FileResponse,
    tags=["transcription"],
    summary="Transcribe Audio File",
    description="Upload an audio file and receive a formatted transcript document",
    responses={
        200: {
            "description": "Transcript generated successfully",
            "content": {
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document": {
                    "schema": {"type": "string", "format": "binary"}
                }
            }
        },
        400: {
            "description": "Invalid file or request parameters",
            "model": ErrorResponse
        },
        413: {
            "description": "File too large",
            "model": ErrorResponse
        },
        422: {
            "description": "Audio processing failed",
            "model": ErrorResponse
        },
        429: {
            "description": "Rate limit exceeded",
            "model": ErrorResponse
        },
        500: {
            "description": "Transcription failed",
            "model": ErrorResponse
        },
        503: {
            "description": "Service unavailable",
            "model": ErrorResponse
        }
    }
)
async def transcribe_audio(
    request: Request,
    file: UploadFile = File(
        ...,
        description="Audio file to transcribe",
        media_type="audio/*"
    ),
    language: str = Query(
        default="en",
        description="Language code for transcription (ISO 639-1)",
        regex="^[a-z]{2}$"
    ),
    settings: Settings = Depends(get_app_settings)
):
    """
    Upload an audio file and receive a formatted transcript document.

    **Supported Audio Formats:**
    - M4A (Apple Audio)
    - MP4 (MPEG-4 Audio)
    - WAV (Waveform Audio)
    - MP3 (MPEG Audio Layer III)
    - FLAC (Free Lossless Audio Codec)
    - OGG (Ogg Vorbis)

    **Features:**
    - Automatic speaker identification
    - Audio preprocessing and optimization
    - Conversation flow analysis
    - Timestamp generation
    - Word document output with formatting

    **Processing Steps:**
    1. File validation and security checks
    2. Audio preprocessing (normalization, silence removal)
    3. Whisper AI transcription
    4. Speaker assignment using conversation patterns
    5. Document generation with timestamps

    **Rate Limits:**
    - Maximum 2 requests per minute per IP
    - File size limit: 100MB
    - Processing timeout: 5 minutes

    **Response:**
    Returns a Microsoft Word document (.docx) containing the formatted transcript
    with speaker identification, timestamps, and conversation statistics.
    """
    request_id = getattr(request.state, 'request_id', str(uuid.uuid4()))
    start_time = time.time()

    # Validate model is loaded
    if not whisper_model:
        logger.error("Whisper model not loaded", extra={"request_id": request_id})
        raise ServiceUnavailableException("Transcription service is not available")

    # Validate file
    try:
        file_info = await validate_uploaded_file(file, settings)
        logger.info(f"File validation passed", extra={
            "request_id": request_id,
            "filename": file_info.filename,
            "size": file_info.size,
            "type": file_info.content_type
        })
    except FileValidationError as e:
        logger.warning(f"File validation failed: {str(e)}", extra={"request_id": request_id})
        raise BadRequestException(str(e))

    # Generate unique paths
    unique_id = str(uuid.uuid4())
    temp_audio_path = os.path.join(settings.upload_dir, f"{unique_id}{file_info.extension}")
    output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.docx').name

    try:
        # Save uploaded file
        logger.info("Saving uploaded file", extra={"request_id": request_id})
        with open(temp_audio_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Preprocess audio
        logger.info("Starting audio preprocessing", extra={"request_id": request_id})
        processed_audio = preprocess_audio(temp_audio_path, settings)

        # Transcribe
        logger.info("Starting transcription", extra={
            "request_id": request_id,
            "language": language,
            "model": settings.whisper_model_name
        })
        segments = transcribe_fast(whisper_model, processed_audio, language, settings)

        # Assign speakers
        logger.info("Assigning speakers", extra={"request_id": request_id})
        speaker_segments = assign_speakers_fast(segments)

        # Create transcript
        logger.info("Creating transcript document", extra={"request_id": request_id})
        create_fast_transcript(speaker_segments, file.filename, output_path)

        # Clean up temp files (but keep output_path for FileResponse)
        cleanup_temp_files([temp_audio_path, processed_audio])

        # Verify output file was created
        if not os.path.exists(output_path):
            logger.error("Transcript file was not created", extra={"request_id": request_id})
            raise InternalServerErrorException("Failed to create transcript file")

        # Generate clean filename
        original_name = Path(file.filename).stem
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        clean_filename = f"transcript_{original_name}_{timestamp}.docx"

        # Log completion
        duration = time.time() - start_time
        logger.info(f"Transcription completed successfully", extra={
            "request_id": request_id,
            "duration": duration,
            "output_file": clean_filename
        })

        # Return file with cleanup
        return FileResponse(
            path=output_path,
            filename=clean_filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            background=BackgroundTask(cleanup_temp_files, [output_path])
        )

    except FileValidationError as e:
        cleanup_temp_files([temp_audio_path])
        logger.warning(f"File validation error: {str(e)}", extra={"request_id": request_id})
        raise BadRequestException(str(e))

    except AudioProcessingError as e:
        cleanup_temp_files([temp_audio_path, processed_audio])
        logger.error(f"Audio processing error: {str(e)}", extra={"request_id": request_id})
        raise BadRequestException(f"Audio processing failed: {str(e)}")

    except TranscriptionError as e:
        cleanup_temp_files([temp_audio_path, processed_audio])
        logger.error(f"Transcription error: {str(e)}", extra={"request_id": request_id})
        raise InternalServerErrorException(f"Transcription failed: {str(e)}")

    except Exception as e:
        cleanup_temp_files([temp_audio_path, processed_audio])
        log_error(logger, request_id, e, "transcription")
        raise InternalServerErrorException("An unexpected error occurred during transcription")

def preprocess_audio(audio_path: str, settings: Settings) -> str:
    """Preprocess audio for optimal Whisper performance"""
    logger.info("Starting audio preprocessing")

    try:
        audio = AudioSegment.from_file(audio_path)
        duration_seconds = len(audio) / 1000
        logger.info(f"Audio duration: {duration_seconds:.1f} seconds")

        # Convert to mono
        if audio.channels > 1:
            audio = audio.set_channels(1)
            logger.info("Converted to mono")

        # Resample to target sample rate
        if audio.frame_rate != settings.target_sample_rate:
            audio = audio.set_frame_rate(settings.target_sample_rate)
            logger.info(f"Resampled to {settings.target_sample_rate}Hz")

        # Normalize and trim silence
        audio = audio.normalize()
        audio = audio.strip_silence(
            silence_len=settings.min_silence_duration,
            silence_thresh=settings.silence_threshold
        )
        logger.info("Applied normalization and silence trimming")

        # Handle long audio
        if duration_seconds > settings.max_audio_length:
            logger.info(f"Audio too long ({duration_seconds:.1f}s), splitting into chunks")
            return split_audio_into_chunks(audio, settings)

        # Save optimized audio
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.wav')
        audio.export(temp_file.name, format='wav')
        logger.info(f"Preprocessed audio saved to {temp_file.name}")
        return temp_file.name

    except Exception as e:
        logger.error(f"Audio preprocessing failed: {str(e)}", exc_info=True)
        raise AudioProcessingError(f"Failed to preprocess audio: {str(e)}")

def split_audio_into_chunks(audio: AudioSegment, settings: Settings) -> list:
    """Split long audio into manageable chunks"""
    chunks = []
    chunk_length_ms = settings.chunk_size * 1000

    for i in range(0, len(audio), chunk_length_ms):
        chunk = audio[i:i + chunk_length_ms]
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.wav')
        chunk.export(temp_file.name, format='wav')
        chunks.append(temp_file.name)

    logger.info(f"Split audio into {len(chunks)} chunks")
    return chunks

def transcribe_fast(model, audio_path, language: str, settings: Settings):
    """Fast transcription with optimized settings"""
    logger.info("Starting transcription")

    config = {
        "language": language,
        "task": "transcribe",
        "verbose": False,
        "condition_on_previous_text": False,
        "no_speech_threshold": 0.6,
        "logprob_threshold": -1.0,
        "temperature": 0.0,
        "compression_ratio_threshold": 2.4,
        "beam_size": 1,
        "best_of": 1,
        "fp16": True,
    }

    if isinstance(audio_path, list):
        return transcribe_chunks(model, audio_path, config, settings)
    else:
        result = model.transcribe(audio_path, **config)
        return process_transcription_result(result, 0)

def transcribe_chunks(model, audio_chunks, config, settings: Settings):
    """Transcribe multiple audio chunks"""
    all_segments = []
    time_offset = 0

    for chunk_idx, chunk_path in enumerate(audio_chunks):
        logger.info(f"Processing chunk {chunk_idx+1}/{len(audio_chunks)}")
        result = model.transcribe(chunk_path, **config)
        segments = process_transcription_result(result, time_offset)
        all_segments.extend(segments)
        time_offset += settings.chunk_size

        try:
            os.unlink(chunk_path)
        except Exception as e:
            logger.warning(f"Failed to delete chunk file {chunk_path}: {str(e)}")

    return all_segments

def process_transcription_result(result, time_offset=0):
    """Process and clean transcription results"""
    segments = result.get("segments", [])
    clean_segments = []
    
    for segment in segments:
        text = clean_transcription_text(segment["text"])
        
        if not is_meaningful_text(text):
            continue
        
        if segment.get("avg_logprob", 0) < -1.5:
            continue
        
        clean_segments.append({
            "start": segment["start"] + time_offset,
            "end": segment["end"] + time_offset,
            "text": text,
            "confidence": segment.get("avg_logprob", 0)
        })
    
    return clean_segments

def clean_transcription_text(text):
    """Clean transcription text"""
    if not text:
        return ""
    
    text = re.sub(r'\b(nd|uh|um|ah)\b', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def is_meaningful_text(text):
    """Check for meaningful content"""
    if not text or len(text.strip()) < 3:
        return False
    
    words = text.split()
    if len(words) < 1:
        return False
    
    text_lower = text.lower()
    artifacts = ['odi recorder', 'install whisper', 'pip install']
    if any(artifact in text_lower for artifact in artifacts):
        return False
    
    return True

def assign_speakers_fast(segments):
    """Improved speaker assignment for natural conversation flow"""
    if not segments:
        return []

    speaker_segments = []
    current_speaker = 1
    last_end_time = 0

    for i, segment in enumerate(segments):
        text = segment["text"].strip().lower()
        pause_duration = segment["start"] - last_end_time

        should_switch = False

        # Long pause suggests speaker change
        if pause_duration > 2.0:
            should_switch = True

        # Response indicators
        response_words = ["yes", "no", "yeah", "okay", "right", "well", "so", "but", "and", "oh"]
        if any(text.startswith(word + " ") or text == word for word in response_words):
            should_switch = True

        # Question-answer pattern
        if i > 0:
            prev_text = segments[i-1]["text"].strip()
            if prev_text.endswith("?"):
                should_switch = True

        if should_switch:
            current_speaker = 2 if current_speaker == 1 else 1

        speaker_segments.append({
            "start": segment["start"],
            "end": segment["end"],
            "text": segment["text"],
            "speaker": f"Speaker {current_speaker}",
            "confidence": segment["confidence"],
            "prev_end": last_end_time
        })

        last_end_time = segment["end"]

    return speaker_segments

def create_fast_transcript(segments, original_filename, output_path):
    """Create transcript with proper conversation flow"""

    doc = Document()
    doc.add_heading("Meeting Transcript", 0)
    doc.add_paragraph(f"Audio File: {original_filename}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("")

    if not segments:
        doc.add_paragraph("No meaningful conversation detected.")
        doc.save(output_path)
        return

    doc.add_heading("Conversation", 1)

    conversation_turns = []
    current_speaker = None
    current_text = ""
    current_start_time = 0

    for seg in segments:
        text = seg["text"].strip()
        speaker = seg["speaker"]

        if speaker != current_speaker:
            if current_speaker and current_text:
                conversation_turns.append({
                    "speaker": current_speaker,
                    "text": current_text.strip(),
                    "start_time": current_start_time
                })

            current_speaker = speaker
            current_text = text
            current_start_time = seg["start"]
        else:
            time_gap = seg["start"] - seg.get("prev_end", seg["start"])

            if time_gap > 5.0 and current_text:
                conversation_turns.append({
                    "speaker": current_speaker,
                    "text": current_text.strip(),
                    "start_time": current_start_time
                })
                current_text = text
                current_start_time = seg["start"]
            else:
                if current_text and not current_text.endswith(('.', '!', '?')):
                    current_text += " " + text
                else:
                    current_text += " " + text

    # Add final turn
    if current_speaker and current_text:
        conversation_turns.append({
            "speaker": current_speaker,
            "text": current_text.strip(),
            "start_time": current_start_time
        })

    # Write conversation
    for turn in conversation_turns:
        text = turn["text"]

        if text and not text[-1] in '.!?':
            text += "."

        if text:
            text = text[0].upper() + text[1:] if len(text) > 1 else text.upper()

        minutes = int(turn["start_time"] // 60)
        seconds = int(turn["start_time"] % 60)
        timestamp = f"[{minutes:02d}:{seconds:02d}]"

        doc.add_paragraph(f"{timestamp} {turn['speaker']}: {text}")

    # Summary
    doc.add_paragraph("")
    doc.add_heading("Summary", 1)

    speaker_stats = {}
    for turn in conversation_turns:
        speaker = turn["speaker"]
        if speaker not in speaker_stats:
            speaker_stats[speaker] = {"turns": 0, "words": 0}
        speaker_stats[speaker]["turns"] += 1
        speaker_stats[speaker]["words"] += len(turn["text"].split())

    doc.add_paragraph(f"Total conversation turns: {len(conversation_turns)}")
    for speaker, stats in speaker_stats.items():
        doc.add_paragraph(f"{speaker}: {stats['turns']} turns, {stats['words']} words")

    doc.save(output_path)

def cleanup_temp_files(file_paths):
    """Clean up temporary files"""
    if isinstance(file_paths, str):
        file_paths = [file_paths]

    for path in file_paths:
        try:
            if path and os.path.exists(path) and ('tmp' in path or 'uploads' in path or path.endswith('.docx')):
                os.unlink(path)
        except:
            pass

# For Vercel deployment
handler = app

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
